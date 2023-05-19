# TODO: add check for "\u2026" in journal field
import json, sys, re, sys

use_docx = True

if use_docx:
    # see https://python-docx.readthedocs.io/en/latest/
    import docx

#temporary!
tmpf = open('tmp.txt', 'w')
tmpf.close()

starting_year = 2015

year_range = None
#year_range = {'after_year':2021, 'before_year':2023}

# Words in journal titles that should be left as-is and not capitalized (or decapitalized)
journal_asis = [
    'and', 
    'in', 
    'on', 
    'of', 
    'the',
    'for',
    'et',
    'de',
    'del',
    'as',
    'la',
    'der',
    'ACS',
    'BioScience',
    'MBio',
    'ISME',
    'PLoS',
    'PeerJ',
    'BMC',
    'della',
    'SORT-Statistics',
    '(Punta',
    'Arenas)',
    'DNA',
    'Asia-Pacific',
    'SSAR',
    'IAWA',
    'USA',
    'KIOES',
    'Socio-Ecological',
    'ZooKeys',
    'F1000Research',
    'GigaScience'
]        

# Authors with surnames identical to EEB faculty members who should be 
# ignored when searching for EEB faculty names in author lists        
authors_like_faculty = [
    'RR Colwell',
    'C Anderson',
    'CD Anderson',
    'DL Anderson',
    'E Anderson',
    'J Anderson',
    'M Anderson',
    'MKJ Anderson',
    'RL Anderson',
    'RM Anderson',
    'T Anderson',
    'TJ Anderson',
    'SE Bush',
    'AM Cooley',
    'JE Cooley',
    'CJ Davis', 
    'D Davis',   
    'DA Davis',   
    'DR Davis', 
    'KE Davis',  
    'JI Davis',
    'JM Davis',
    'JP Davis',
    'WJ Davis',
    'N Henry',
    'WJ Henry',
    'AT Jones',
    'AW Jones',
    'FA Jones',
    'FAM Jones',
    'TH Jones',
    'AM Les',
    'FH Wagner',
    'DM Wagner',
    'M Wagner',
    'V Wagner',
    'CB Schultz',
    'A Simon',
    'MF Simon',
    'B Wells',
    'Z Yuan'
]

singleton = None #['Bagchi']
faculty = [
    'Anderson',
    'Bagchi',
    'Bolnick',
    'Bush',
    'Caira',
    'Chazdon',
    'Colwell',
    'Cooley',
    'Davis',
    'Diggle',
    'Elphick',
    'Fusco',
    'GarciaRobledo',
    'Goffinet',
    'Henry',
    'Herrick',
    'Holsinger',
    'Jockusch',
    'Jones',
    'Knutie',
    'LewisL',
    'LewisP',
    'Les',
    'Likens',
    'Schlichting',
    'Schultz',
    'Schwenk',
    'Seemann',
    'Silander',
    'Simon',
    'Tingley',
    'Trumbo',
    'Turchin',
    'Urban',
    'Wagner',
    'Wegrzyn',
    'Wells',
    'Willig',
    'Yarish',
    'Yuan'
]

filepaths = {
    'Anderson':      'Gregory-Anderson-final.json',
    'Bagchi':        'Robert-Bagchi-final.json',
    'Bolnick':       'Daniel-Bolnick-final.json',
    'Bush':          'Andrew-Bush-final.json',
    'Caira':         'Janine-Caira-final.json',
    'Chazdon':       'Robin-Chazdon-final.json',
    'Colwell':       'Rob-Colwell-final.json',
    'Cooley':        'John-Cooley-final.json',
    'Davis':         'Miranda-Davis-final.json',
    'Diggle':        'Pamela-Diggle-final.json',
    'Elphick':       'Chris-Elphick-final.json',
    'Fusco':         'Nicole-Fusco-final.json',
    'GarciaRobledo': 'Carlos-GarciaRobledo-final.json',
    'Goffinet':      'Bernard-Goffinet-final.json',
    'Henry':         'Charles-Henry-final.json',
    'Herrick':       'Susan-Herrick-final.json',
    'Holsinger':     'Kent-Holsinger-final.json',
    'Jockusch':      'Elizabeth-Jockusch-final.json',
    'Jones':         'Cynthia-Jones-final.json',
    'Knutie':        'Sarah-Knutie-final.json',
    'LewisL':        'Louise-Lewis-final.json',
    'LewisP':        'Paul-Lewis-final.json',
    'Les':           'Donald-Les-final.json',
    'Likens':        'Gene-Likens-final.json',
    'Schlichting':   'Carl-Schlichting-final.json',
    'Schultz':       'Eric-Schultz-final.json',
    'Schwenk':       'Kurt-Schwenk-final.json',
    'Seemann':       'Jeff-Seemann-final.json',
    'Silander':      'John-Silander-final.json',
    'Simon':         'Chris-Simon-final.json',
    'Tingley':       'Morgan-Tingley-final.json',
    'Trumbo':        'Stephen-Trumbo-final.json',
    'Turchin':       'Peter-Turchin-final.json',
    'Urban':         'Mark-Urban-final.json',
    'Wagner':        'David-Wagner-final.json',
    'Wegrzyn':       'Jill-Wegrzyn-final.json',
    'Wells':         'Kentwood-Wells-final.json',
    'Willig':        'Michael-Willig-final.json',
    'Yarish':        'Charles-Yarish-final.json',
    'Yuan':          'Yaowu-Yuan-final.json'
}

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """
    if use_docx:
        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if not color is None:
          c = docx.oxml.shared.OxmlElement('w:color')
          c.set(docx.oxml.shared.qn('w:val'), color)
          rPr.append(c)

        # Remove underlining if it is requested
        if not underline:
          u = docx.oxml.shared.OxmlElement('w:u')
          u.set(docx.oxml.shared.qn('w:val'), 'none')
          rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)
    else:
        hyperlink = '%s (%s)' % (text, url)
        
    return hyperlink

def getNumCitations(result):
    ncites = 0
    try:
        ncites = result['ncites']
    except KeyError:
        print(result)
        sys.exit('No key named "ncites"');
    if ncites is not None and not ncites.__class__.__name__ == 'int':
        print(result)
        sys.exit('"ncites" is not an integer');
    if ncites is None:
        ncites  = 0
    return ncites

# Assumes ncites is not None and assumes either result['authors'] or result['bookeditors'] is not None. 
# Returns dictionary containing:
# 'authors' (string): formatted authors list (<eeb author name> converted to **eeb author name**)
# 'surnames' (string): list of just surnames, used to alphabetize entries
# 'eeb_cite_counts' (vector): each element is a tuple (EEB author surname, citation count)
def updateCounts(result, ncites):
    assert ncites is  not None, 'ncites passed to updateCounts was None'
    
    #global person_counts, cites_vect
    authors = result['authors']
    if authors is None:
        authors = result['bookeditors']
    assert authors is not None, "Both result['authors'] and result['bookeditors'] passed to updateCounts was None"
        
    # Split the authors list at the commas
    author_list = authors.split(',')
    
    # This will hold number of EEB authors on this particular publication
    num_eeb_authors = 0
    
    # This will be a vector of surnames used for sorting bibliographic entries alphabetically
    # It will be returned concatenated into a string under the key 'surnames' in the returned dictionary
    surname_vect = []
    
    # This will be returned under the key 'eeb_cite_counts' in the returned dictionary
    eeb_cite_counts = []
    
    # This will be returned concatenated into a string under the key 'authors' in the returned dictionary
    processed_author_vect = []
    
    # Process each author in author_list
    for a in author_list:
        astripped = a.strip();
        
        # If author is marked as an EEB author, remove brackets around name
        # (the fact that m is not None tells us later that this author was an EEB author)
        m = re.match('<(.+)>', astripped)
        if m is not None:
            astripped = m.group(1)
        
        # Handle papers involving too many authors to list
        # e.g. <...G Likens...15364>
        mm = re.match('[.][.][.](.+?)[.][.][.](\d+)', astripped)
        num_others = 0
        if mm is not None:
            astripped = mm.group(1)
            num_others = int(mm.group(2))
            #input('astripped = "%s", num_others = %d' % (astripped, num_others))
        
        # Split author name at spaces
        eeb_parts = astripped.split()
        
        # Separate author name into initials and surname
        # There are a lot of special cases in which one of the partial_surnames listed
        # below should be considered as part of the surname
        initials = None
        surname = None
        partial_surnames = ['al', 'Al', 'dalla', 'Dalla', 'da', 'Da', 'de', 'De', 'del', 'Del', 'den', 'Den', 'des', 'Des', 'di', 'Di', 'do', 'Do', 'dos', 'Dos', 'la', 'La', 'le', 'Le', 'san', 'San', 'st', 'St', 'ter', 'Ter', 'van', 'Van', 'von', 'Von']
        if len(eeb_parts) == 5:
            ok = False
            if eeb_parts[1] in partial_surnames:
                initials = eeb_parts[0]
                surname = '%s %s %s %s' % (eeb_parts[1], eeb_parts[2], eeb_parts[3], eeb_parts[4])
                ok = True
            elif astripped == '(511 authors including <MR Willig>)':
                initials = ''
                surname = astripped
                ok = True
            assert ok, 'eeb_parts has length 5 but does not fall into acceptable cases: author = "%s"' % astripped
        elif len(eeb_parts) == 4:
            ok = False
            if eeb_parts[1] in partial_surnames:
                initials = eeb_parts[0]
                surname = '%s %s %s' % (eeb_parts[1], eeb_parts[2], eeb_parts[3])
                ok = True
            assert ok, 'eeb_parts has length 4 but does not fall into acceptable cases: author = "%s"' % astripped
        elif len(eeb_parts) == 3:
            ok = False
            if eeb_parts[2] in ['Jr', 'II', 'III']:
                initials = eeb_parts[0]
                surname = '%s %s' % (eeb_parts[1], eeb_parts[2])
                ok = True
            elif eeb_parts[1] in partial_surnames:
                initials = eeb_parts[0]
                surname = '%s %s' % (eeb_parts[1], eeb_parts[2])
                ok = True
            elif 'D Ortega‐Del Vecchyo' == astripped:
                initials = 'D'
                surname = 'Ortega‐Del Vecchyo'
                ok = True
            assert ok, 'eeb_parts has length 3 but does not fall into acceptable cases: author = "%s" (eeb_parts[2] = "%s")' % (astripped, eeb_parts[2])
        elif len(eeb_parts) == 2:
            initials = eeb_parts[0]
            surname = eeb_parts[1]
        elif len(eeb_parts) == 1:
            initials = ''
            surname = eeb_parts[0]
        elif 'et al.' in astripped:
            initials = ''
            surname = astripped
        else:
            assert False, 'Could not parse EEB author "%s"' % astripped

        # Add surname to vector of surnames
        surname_vect.append(surname)
        
        # Add author name to vector of formatted author names
        processed_author = '%s %s' % (initials, surname)
        if m is not None:
            # Bold because this is an EEB author
            processed_author = '**%s %s**' % (initials, surname)
        if num_others > 0:
            processed_author = '(%d others including %s)' % (num_others, processed_author)
        processed_author_vect.append(processed_author)
        
        # Check to make sure authors NOT marked as EEB faculty members
        # are indeed not EEB faculty members
        if m is None:
            if surname in faculty and not astripped in authors_like_faculty:
                print(result)
                sys.exit('author "%s" needs to be bracketed' % astripped)

        # Add citations to returned eeb_cite_counts vector
        if m is not None:
            num_eeb_authors += 1
            if surname == 'Lewis':
                if initials == 'L' or initials == 'LA':
                    surname = 'LLewis'
                elif initials == 'P' or initials == 'PO':
                    surname = 'PLewis'
        
            eeb_cite_counts.append((surname, ncites))
       
    # Concatenate surnames into a surnames_str that is returned          
    try:
        surnames_str = ','.join(surname_vect)
    except TypeError:
        print(report)
        sys.exit('Could not join surnames of authors into a string')
        
    # Concatenate processed_authors_list into a processed_author_str that is returned
    processed_author_str = ', '.join(processed_author_vect)

    # Check to ensure that every publication has at least one EEB author        
    if num_eeb_authors == 0:
        print('num_eeb_authors == 0:')
        print('  authors: %s' % authors)
        print('  surnames_str: %s' % surnames_str)
        print(result)
        sys.exit('This publication has no EEB authors')
                
    return {'authors':processed_author_str, 'surnames':surnames_str, 'eeb_cite_counts':eeb_cite_counts}
    
def checkJournalNames(journalf, journal):
    if journal is None:
        return None
        
    diff = ord('a') - ord('A')
    ordA = ord('A')
    ordZ = ord('Z')

    # Create new journal name with every word capitalized
    jvect = []
    jparts = journal.strip().split(' ')
    for jpart in jparts:
        comma = False
        m = re.match('(.+),', jpart)
        if m is not None:
            comma = True
            jpart = m.group(1)
        if jpart == '&':
            jpart = 'and'
        elif not jpart in journal_asis:
            jpart = jpart.capitalize()
        if comma:
            jpart += ','
        jvect.append(jpart)
        
    journal2 = ' '.join(jvect)
    ok = journal == journal2
        
    if not ok:
        journalf.write('"%s" --> "%s"\n' % (journal, journal2))
        return journal2
        
    return journal

def dumpException(exceptf, f, a, y, t, j, v, b, e, doi):
    exceptf.write('-------- %s ---------\n' % f)

    if a:
        exceptf.write('authors: %s\n' % a)
    else:
        exceptf.write('authors: MISSING\n')

    if y:
        exceptf.write('year: %s\n' % y)
    else:
        exceptf.write('year: MISSING\n')

    if t:
        exceptf.write('title: %s\n' % t)
    else:
        exceptf.write('title: MISSING\n')

    if j:
        exceptf.write('journal: %s\n' % j)
    else:
        exceptf.write('journal: MISSING\n')

    if v:
        exceptf.write('volume: %s\n' % v)
    else:
        exceptf.write('volume: MISSING\n')

    if b:
        exceptf.write('bpage: %s\n' % b)
    else:
        exceptf.write('bpage: MISSING\n')

    if e:
        exceptf.write('epage: %s\n' % e)
    else:
        exceptf.write('epage: MISSING\n')

    if doi:
        exceptf.write('https://doi.org/%s\n' % doi)
    else:
        exceptf.write('doi: MISSING\n')

exceptf = open('exceptions.txt', 'w')
journalf = open('journal-issues.txt', 'w')
dupf = open('duplicates.txt', 'w')
nyearless = 0
nexceptions = 0
nduplicates = 0
ninpress = 0
narticles = 0
nbooks = 0
nchapters = 0
ncites_dept = 0
neditedvolumes = 0
bibentries = []
titles_seen = {}
title_lookup = {}
person_counts = {}
cites_vect = []

chosen_ones = singleton
if chosen_ones is None:
    chosen_ones = faculty
for f in chosen_ones:
    fn = '%s' % filepaths[f]
    print('Reading file "%s"...' % fn)
    stuff = open(fn, 'r').read()
    results = json.loads(stuff)
    
    for result in results:
        if not result['ignore']:
            # Add to counts for EEB authors
            has_authors = result['authors'] is not None
            
            ncites = getNumCitations(result)
            update_results = updateCounts(result,ncites)
            authors = update_results['authors']
            surnames = update_results['surnames']

            year    = result['year']
            title   = result['title']

            journal = result['journal']
            good_journal_name = checkJournalNames(journalf, journal)
            if not good_journal_name == journal:
                print(result)
                sys.exit('bad journal name: "%s" --> "%s"' % (journal,good_journal_name))

            volume  = result['volume']
            number  = result['number']
            
            # See if this is an edited volume
            try:
                booktitle      = result['booktitle']
                bookeditors    = result['bookeditors']
                bookpublisher  = result['bookpublisher']
                bookcity       = result['bookcity']
                bookisbn       = result['bookisbn']
            except KeyError:
                booktitle      = None
                bookeditors    = None
                bookpublisher  = None
                bookcity       = None
                bookisbn       = None
                
            # If edited volume in which one of the editors is EEB faculty, and 
            # if result['authors'] is None, replace bookeditors with authors 
            # processed by updateCounts
            if not has_authors:
                if bookeditors is not None:
                    bookeditors = authors
                else:
                    print(result)
                    sys.exit('authors for above paper was None')
            
            bpage   = result['bpage']
            epage   = result['epage']
            article = result['article']
            doi     = result['doi']
            url     = result['url']
            citation_id = result['citation_id']
                            
            # Construct bibliography string
            yearless = False
            exception = False
            inpress = False
            ischapter = False
            iseditedvolume = False
            isbook = False
            bib = ''
            if has_authors:
                bib += '%s.' % authors
            if year:
                bib += ' %s.' % year
            else:
                yearless = True
            if title:
                bib += ' %s.' % title
            if journal and volume and bpage and epage:
                bib += ' %s %s:%s-%s.' % (journal, volume, bpage, epage)
            elif journal and volume and article:
                bib += ' %s %s:%s.' % (journal, volume, article)
            elif journal and volume == 'inpress':
                inpress = True
                bib += ' %s (in press).' % journal
            elif (not has_authors) and (booktitle and bookeditors and bookpublisher and bookcity and bpage and epage):
                iseditedvolume = True
                bib = '%s (eds.) %s. %s. %s, %s.' % (bookeditors, year, booktitle, bookpublisher, bookcity)
            elif booktitle and bookeditors and bookpublisher and bookcity and bpage and epage:
                ischapter = True
                bib += ' pp. %s-%s in: %s, %s (eds.) %s, %s.' % (bpage, epage, booktitle, bookeditors, bookpublisher, bookcity)
            elif booktitle and bookeditors and bookpublisher and bpage and epage:
                ischapter = True
                bib += ' pp. %s-%s in: %s, %s (eds.) %s.' % (bpage, epage, booktitle, bookeditors, bookpublisher)
            elif booktitle and bookpublisher:
                ischapter = True
                bib += ' In: %s, %s' % (booktitle, bookpublisher)
            elif bookisbn and bookpublisher:
                isbook = True
                bib += ' %s' % bookpublisher
                if bookcity:
                    bib += ', %s' % bookcity
                else:
                    bib += '.'
            else:
                exception = True
                
            # Add DOI or ISBN if available
            if doi:
                if use_docx:
                    bib += ' <<<https://doi.org/%s>>>' % doi
                else:
                    bib += ' [https://doi.org/%s](https://doi.org/%s)' % (doi,doi)
            elif bookisbn:
                bib += ' ISBN: %s.' % bookisbn
                
            # Add number of citations
            if ncites == 1:
                bib += ' (1 citation)'
            else:
                bib += ' (%d citations)' % ncites
            bib += '\n'
            
            if not has_authors:
                nexceptions += 1
                dumpException(exceptf, f, 'NO AUTHORS!', year, title, journal, volume, bpage, epage, doi)
                
            if yearless or (year_range is None) or (year > year_range['after_year'] and year < year_range['before_year']):
                already_seen = False
                if title in titles_seen.keys():
                    already_seen = True
                else:
                    titles_seen[title] = f
                
                # Update citation counts
                if ncites is not None and not already_seen:
                    ncites_dept += ncites
                cite_counts = update_results['eeb_cite_counts']
                for entry in cite_counts:
                    surname = entry[0]
                    ncites = entry[1]
                    
                    assert surname is not None
                    assert ncites.__class__.__name__ == 'int', 'ncites is a %s' % ncites.__class__.__name__
                    
                    # Update person_counts, incrementing number of papers ('works') 
                    # and number of citations ('cites')
                    if surname in person_counts.keys():
                        person_counts[surname]['works'] += 1
                        person_counts[surname]['cites'] += ncites
                    else:
                        person_counts[surname] = {'works':1, 'cites':ncites}

                # Update cites_vect, adding tuple (num EEB authors, num citations)
                cites_vect.append((len(cite_counts),ncites))
                
                #temporary!
                #tmpf = open('tmp.txt', 'a')
                #tmpf.write()
                #tmpf.close()
                    
                # Save in bibentries list
                if yearless:
                    nyearless += 1
                    nexceptions += 1
                    dumpException(exceptf, f, authors, year, title, journal, volume, bpage, epage, doi)
                elif exception:
                    nexceptions += 1
                    dumpException(exceptf, f, authors, year, title, journal, volume, bpage, epage, doi)
                elif ischapter:
                    entry = (year, surnames, bib)
                    if already_seen:
                        nduplicates += 1
                        dupf.write('\n-------------------\n')
                        dupf.write('Previous: %s\n' % titles_seen[title])
                        dupf.write('~~> %s\n' % bib)
                        dupf.write('Current: %s\n' % f)
                        dupf.write('~~> %s\n' % title_lookup[title])
                    else:
                        nchapters += 1
                        bibentries.append(entry)
                        title_lookup[title] = bib
                elif isbook:
                    nbooks += 1
                    entry = (year, surnames, bib)
                    bibentries.append(entry)
                elif iseditedvolume:
                    entry = (year, surnames, bib)
                    if already_seen:
                        nduplicates += 1
                        dupf.write('\n-------------------\n')
                        dupf.write('Previous: %s\n' % titles_seen[title])
                        dupf.write('~~> %s\n' % bib)
                        dupf.write('Current: %s\n' % f)
                        dupf.write('~~> %s\n' % title_lookup[title])
                    else:
                        neditedvolumes += 1
                        bibentries.append(entry)
                        title_lookup[title] = bib
                else:
                    entry = (year, surnames, bib)
                    if already_seen:
                        nduplicates += 1
                        dupf.write('\n-------------------\n')
                        dupf.write('Previous: %s\n' % titles_seen[title])
                        dupf.write('~~> %s\n' % bib)
                        dupf.write('Current: %s\n' % f)
                        dupf.write('~~> %s\n' % title_lookup[title])
                    else:
                        narticles += 1
                        bibentries.append(entry)
                        title_lookup[title] = bib
                if inpress:
                    ninpress += 1
                    
exceptf.close()
dupf.close()
journalf.close()

bibentries.sort()
ngood = len(bibentries)

if use_docx:
    emu = 914400 # equals 1 inch  (not used currently, but useful for specifying widths of table cells)
    doc = docx.Document()
    for i,b in enumerate(bibentries):
        bibentry = b[2]
        parts = re.split('[*][*](.+?)[*][*]', bibentry)
        para = doc.add_paragraph()
        para.add_run('%d. ' % (i+1,))
        for i,p in enumerate(parts):
            if i % 2 == 0:
                subparts = re.split('<<<(.+?)>>>', p)
                if len(subparts) == 3:
                    para.add_run(subparts[0])
                    url = subparts[1]
                    add_hyperlink(para, url, url, '0000FF', False)
                    para.add_run(subparts[2])
                else:
                    para.add_run(p)
            else:
                para.add_run(p).bold = True
    doc.save('bibliography.docx')
else:
    gatherf = open('bibliography.md', 'w')
    gatherf.write('\nSorted bibliographic entires:\n')
    for i,b in enumerate(bibentries):
        gatherf.write('%d. %s\n' % (i+1, b[2]))
    gatherf.close()

print('nyearless      = %d' % nyearless)
print('nexceptions    = %d' % nexceptions)
print('nduplicates    = %d' % nduplicates)
print('ninpress       = %d' % ninpress)
print('narticles      = %d' % narticles)
print('nchapters      = %d' % nchapters)
print('nbooks         = %d' % nbooks)
print('neditedvolumes = %d' % neditedvolumes)
print('ngood          = %d' % ngood)
print('narticles + nchapters + nbooks + neditedvolumes= %d' % (narticles + nchapters + nbooks + neditedvolumes,))

print('\nPerson counts:')
counts = []
for k in person_counts.keys():
    counts.append((person_counts[k]['works'], k))
counts.sort()
counts.reverse()
for c,p in counts:
    print('%6d %s' % (c,p))

print('\nCitation counts:')
counts = []
for k in person_counts.keys():
    counts.append((person_counts[k]['cites'], k))
counts.sort()
counts.reverse()
ncites_cum = 0.0
for c,p in counts:
    print('%6d %s' % (c,p))
    ncites_cum += c
print('Total number of citations (not corrected for overcounting): %d' % ncites_cum)
print('Total number of citations (avoiding overcounting): %d' % ncites_dept)

print('\nCitations/paper:')
counts = []
for k in person_counts.keys():
    counts.append((float(person_counts[k]['cites'])/person_counts[k]['works'], k))
counts.sort()
counts.reverse()
for c,p in counts:
    print('%6d %s' % (c,p))

print('\nCitations check:')
cites_vect_length = len(cites_vect)
print('%12d cites_vect_length' % cites_vect_length)
print('%12d ngood + nduplicates' % (ngood + nduplicates,))
#total_cites = 0
total_cites_x_eeb_authors = 0
for x in cites_vect:
    #total_cites += x[1]
    total_cites_x_eeb_authors += x[0]*x[1]
#print('total_cites: %d' % total_cites)  # this number is not relevant
print('%12d total_cites_x_eeb_authors' % total_cites_x_eeb_authors)
print('%12d Total number of citations (not corrected for overcounting)' % ncites_cum)
    
